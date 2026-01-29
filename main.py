import sys
import os
import json
from PySide6.QtWidgets import (QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
                             QLabel, QRadioButton, QPushButton, QButtonGroup, 
                             QMessageBox, QScrollArea, QPlainTextEdit, QComboBox, QDialog, QLineEdit, QFrame)
from PySide6.QtCore import QTimer, Qt
from PySide6.QtGui import QFont, QPixmap, QIcon
from docx import Document

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except Exception:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

class CodeEditor(QPlainTextEdit):
    def __init__(self, placeholder="# Viết lời giải tại đây..."):
        super().__init__()
        self.setFont(QFont("Courier New", 11))
        self.setPlaceholderText(placeholder)
        self.setStyleSheet("""
            QPlainTextEdit {
                background-color: #282c34; 
                color: #abb2bf; 
                border: 1px solid #dcdde1;
                border-radius: 8px; 
                padding: 10px;
                margin-bottom: 20px;
            }
        """)
        self.setMinimumHeight(300)

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key_Return, Qt.Key_Enter):
            cursor = self.textCursor()
            line_text = cursor.block().text()
            indent = len(line_text) - len(line_text.lstrip())
            if line_text.strip().endswith(':'):
                indent += 4
            super().keyPressEvent(event)
            self.insertPlainText(" " * indent)
        elif event.key() == Qt.Key_Tab:
            self.insertPlainText("    ")
        else:
            super().keyPressEvent(event)

class NameInputDialog(QDialog):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Xác nhận thông tin học sinh")
        self.setFixedSize(450, 250)
        self.setStyleSheet("background-color: white;")
        layout = QVBoxLayout(self)
        layout.setContentsMargins(30, 30, 30, 30)
        
        label = QLabel("HỌ VÀ TÊN HỌC SINH:")
        label.setStyleSheet("font-size: 14px; font-weight: bold; color: #2f3640;")
        layout.addWidget(label)
        
        self.name_edit = QLineEdit()
        self.name_edit.setPlaceholderText("Nhập đầy đủ họ tên của em...")
        self.name_edit.setStyleSheet("""
            QLineEdit {
                padding: 12px; 
                font-size: 15px; 
                border: 2px solid #3498db; 
                border-radius: 8px; 
                color: #2f3640;
                background-color: #ffffff;
            }
        """)
        layout.addWidget(self.name_edit)
        
        layout.addSpacing(10)
        self.btn_confirm = QPushButton("NỘP BÀI NGAY")
        self.btn_confirm.setCursor(Qt.PointingHandCursor)
        self.btn_confirm.setStyleSheet("""
            QPushButton {
                background-color: #27ae60; 
                color: white; 
                padding: 12px; 
                font-size: 14px;
                font-weight: bold; 
                border-radius: 8px;
            }
            QPushButton:hover { background-color: #2ecc71; }
        """)
        self.btn_confirm.clicked.connect(self.accept)
        layout.addWidget(self.btn_confirm)

    def get_name(self):
        return self.name_edit.text().strip()

class QuizApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("Hệ thống Kiểm tra Lập trình")
        self.resize(1200, 750)
        self.data_folder = "data"
        self.is_submitted = False
        self.student_name = ""
        
        self.apply_global_styles()
        self.show_selection_screen()
        self.center_on_screen()

    def apply_global_styles(self):
        self.setStyleSheet("""
            QMainWindow { background-color: #f5f6fa; }
            QLabel { color: #2f3640; font-family: 'Segoe UI'; }
            QWidget#content_container { background-color: white; border-radius: 12px; border: 1px solid #dcdde1; }
            QRadioButton { color: #353b48; padding: 8px; font-size: 13px; }
            QRadioButton:hover { background-color: #f1f2f6; border-radius: 5px; }
            QPushButton { border-radius: 8px; color: white; font-weight: bold; font-family: 'Segoe UI'; }
            QComboBox { 
                background-color: white; 
                color: #2f3640; 
                border: 1px solid #dcdde1; 
                padding: 8px; 
                border-radius: 6px;
                font-size: 14px;
            }
            QScrollArea { border: none; background-color: transparent; }
            QMessageBox { background-color: #ffffff; }
            QMessageBox QLabel { color: #2f3640; font-size: 14px; }
            QMessageBox QPushButton { background-color: #3498db; color: white; padding: 6px 15px; min-width: 80px; }
        """)

    def center_on_screen(self):
        screen = QApplication.primaryScreen().availableGeometry()
        size = self.frameGeometry()
        self.move((screen.width() - size.width()) // 2, (screen.height() - size.height()) // 2)

    def show_selection_screen(self):
        self.is_submitted = False
        self.student_name = ""
        self.selection_widget = QWidget()
        self.setCentralWidget(self.selection_widget)
        layout = QVBoxLayout(self.selection_widget)
        layout.setAlignment(Qt.AlignCenter)
        layout.setSpacing(15)

        logo_label = QLabel()
        logo_pixmap = QPixmap(resource_path("logo.png"))
        if not logo_pixmap.isNull():
            logo_label.setPixmap(logo_pixmap.scaled(200, 200, Qt.KeepAspectRatio, Qt.SmoothTransformation))
        else:
            logo_label.setText("🏫")
            logo_label.setStyleSheet("font-size: 80px;")
        layout.addWidget(logo_label, alignment=Qt.AlignCenter)

        title = QLabel("HỆ THỐNG BÀI TẬP")
        title.setStyleSheet("font-size: 28px; font-weight: bold; color: #2980b9; margin-top: 10px;")
        layout.addWidget(title, alignment=Qt.AlignCenter)

        self.combo_modules = QComboBox()
        self.combo_modules.setFixedWidth(450)
        self.module_mapping = {}
        if os.path.exists(self.data_folder):
            files = [f for f in os.listdir(self.data_folder) if f.endswith(".json")]
            for f in files:
                pretty_name = f.replace(".json", "").replace("module", "Module ").capitalize()
                display_name = f"📝 Bài tập {pretty_name}"
                self.module_mapping[display_name] = f
                self.combo_modules.addItem(display_name)
        
        if self.combo_modules.count() == 0:
            self.combo_modules.addItem("Chưa có dữ liệu bài tập (.json)")
        
        layout.addWidget(self.combo_modules, alignment=Qt.AlignCenter)
        layout.addSpacing(20)
        btn_start = QPushButton("BẮT ĐẦU LÀM BÀI")
        btn_start.setFixedSize(250, 55)
        btn_start.setStyleSheet("""
            QPushButton { background-color: #3498db; font-size: 16px; border-bottom: 4px solid #2980b9; }
            QPushButton:hover { background-color: #2980b9; }
            QPushButton:pressed { border-bottom: 0px; margin-top: 4px; }
        """)
        btn_start.clicked.connect(self.start_quiz_logic)
        layout.addWidget(btn_start, alignment=Qt.AlignCenter)

    def start_quiz_logic(self):
        display_name = self.combo_modules.currentText()
        filename = self.module_mapping.get(display_name)
        if not filename: return
        path = os.path.join(self.data_folder, filename)
        try:
            with open(path, "r", encoding="utf-8") as f:
                self.quiz_data = json.load(f)
            self.init_quiz_form()
        except: QMessageBox.critical(self, "Lỗi", "Không thể khởi tạo bài tập!")

    def init_quiz_form(self):
        self.quiz_title = self.quiz_data.get("title", "BaiTap")
        self.questions = self.quiz_data.get("questions", [])
        tasks = self.quiz_data.get("coding_tasks", self.quiz_data.get("coding_task", []))
        self.coding_tasks = tasks if isinstance(tasks, list) else [tasks]
        self.time_left = self.quiz_data.get("time_seconds", 900)
        
        quiz_widget = QWidget()
        self.setCentralWidget(quiz_widget)
        main_layout = QVBoxLayout(quiz_widget)
        main_layout.setContentsMargins(20, 10, 20, 20)

        timer_container = QWidget()
        timer_container.setStyleSheet("background-color: white; border-radius: 10px; border: 1px solid #dcdde1;")
        timer_lay = QHBoxLayout(timer_container)
        self.label_timer = QLabel(f"⏱ THỜI GIAN CÒN LẠI: {self.format_time(self.time_left)}")
        self.label_timer.setAlignment(Qt.AlignCenter)
        self.label_timer.setStyleSheet("font-size: 20px; font-weight: bold; color: #c0392b; border: none;")
        timer_lay.addWidget(self.label_timer)
        main_layout.addWidget(timer_container)

        scroll = QScrollArea()
        scroll.setWidgetResizable(True)
        content_container = QWidget()
        content_container.setObjectName("content_container")
        self.content_layout = QVBoxLayout(content_container)
        self.content_layout.setContentsMargins(30, 20, 30, 20)

        title_q = QLabel("I. PHẦN TRẮC NGHIỆM")
        title_q.setStyleSheet("font-size: 18px; font-weight: bold; color: #2980b9; margin-bottom: 10px;")
        self.content_layout.addWidget(title_q)
        
        self.answer_groups = []
        for i, q in enumerate(self.questions):
            q_box = QWidget()
            q_lay = QVBoxLayout(q_box)
            lbl = QLabel(f"Câu {i+1}: {q['q']}")
            lbl.setWordWrap(True)
            lbl.setStyleSheet("font-weight: bold; font-size: 14px; color: #2f3640;")
            q_lay.addWidget(lbl)
            
            group = QButtonGroup(self)
            self.answer_groups.append(group)
            for idx, opt in enumerate(q['opts']):
                rb = QRadioButton(opt)
                group.addButton(rb, idx)
                q_lay.addWidget(rb)
            self.content_layout.addWidget(q_box)
            line = QFrame(); line.setFrameShape(QFrame.HLine); line.setStyleSheet("color: #f1f2f6;"); self.content_layout.addWidget(line)

        title_c = QLabel("II. PHẦN TỰ LUẬN (CODE)")
        title_c.setStyleSheet("font-size: 18px; font-weight: bold; color: #2980b9; margin-top: 20px; margin-bottom: 10px;")
        self.content_layout.addWidget(title_c)
        
        self.coding_editors = []
        for i, task in enumerate(self.coding_tasks):
            self.content_layout.addWidget(QLabel(f"<b>Bài tập {i+1}:</b> <i style='color: #7f8c8d;'>{task}</i>"))
            editor = CodeEditor(f"# Lời giải cho bài tập {i+1}...")
            self.coding_editors.append(editor)
            self.content_layout.addWidget(editor)

        scroll.setWidget(content_container)
        main_layout.addWidget(scroll)

        btn_layout = QHBoxLayout()
        self.btn_submit = QPushButton("NỘP BÀI")
        self.btn_submit.setStyleSheet("height: 55px; background-color: #27ae60; font-size: 16px;")
        self.btn_submit.clicked.connect(self.confirm_submission)
        
        self.btn_export = QPushButton("XUẤT BẢN SAO LƯU (WORD)")
        self.btn_export.setEnabled(False)
        self.btn_export.setStyleSheet("height: 55px; background-color: #95a5a6; font-size: 15px;")
        self.btn_export.clicked.connect(lambda: self.generate_document(save_locally=True))
        
        btn_layout.addWidget(self.btn_submit, 2)
        btn_layout.addWidget(self.btn_export, 1)
        main_layout.addLayout(btn_layout)

        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_timer)
        self.timer.start(1000)

    def format_time(self, s):
        return f"{s//60:02d}:{s%60:02d}"

    def update_timer(self):
        if self.time_left > 0:
            self.time_left -= 1
            self.label_timer.setText(f"⏱ THỜI GIAN CÒN LẠI: {self.format_time(self.time_left)}")
            if self.time_left <= 60: self.label_timer.setStyleSheet("font-size: 20px; font-weight: bold; color: red; background: #ffeaa7;")
        else: self.submit_process(auto=True)

    def confirm_submission(self):
        if self.is_submitted:
            self.show_selection_screen()
            return

        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Xác nhận")
        msg_box.setIcon(QMessageBox.Question)
        msg_box.setText("<b style='font-size: 16px; color: #2c3e50;'>XÁC NHẬN NỘP BÀI KIỂM TRA?</b>")
        msg_box.setInformativeText("Sau khi nhấn 'Có', em sẽ không thể sửa bài làm nữa.")
        confirm_button = msg_box.addButton("Có, nộp bài", QMessageBox.YesRole)
        cancel_button = msg_box.addButton("Quay lại", QMessageBox.NoRole)
        confirm_button.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold;")
        cancel_button.setStyleSheet("background-color: #e74c3c; color: white; font-weight: bold;")
        msg_box.exec()
        if msg_box.clickedButton() == confirm_button: 
            self.submit_process()

    def submit_process(self, auto=False):
        if self.is_submitted: return

        # Nếu hết giờ (auto=True), bắt buộc phải nhập tên, không cho Cancel
        if auto:
            self.student_name = ""
            while not self.student_name:
                dialog = NameInputDialog(self)
                dialog.setWindowTitle("HẾT GIỜ! VUI LÒNG NHẬP TÊN")
                # Ẩn nút X đóng cửa sổ để ép nhập tên (tùy hệ điều hành)
                dialog.setWindowFlags(dialog.windowFlags() | Qt.CustomizeWindowHint)
                dialog.setWindowFlags(dialog.windowFlags() & ~Qt.WindowCloseButtonHint)
                
                if dialog.exec() == QDialog.Accepted:
                    self.student_name = dialog.get_name()
                    if not self.student_name:
                        QMessageBox.warning(self, "Chú ý", "Em không được để trống họ tên!")
                else:
                    # Nếu học sinh nhấn Esc hoặc tìm cách thoát, thông báo bắt buộc
                    QMessageBox.warning(self, "Bắt buộc", "Hết giờ làm bài, em phải nhập tên để nộp bài!")
        else:
            # Nếu chủ động nộp bài
            dialog = NameInputDialog(self)
            if dialog.exec() != QDialog.Accepted: return
            self.student_name = dialog.get_name()
            if not self.student_name: self.student_name = "Học sinh ẩn danh"

        self.is_submitted = True
        self.timer.stop()
        
        for group in self.answer_groups:
            for btn in group.buttons(): btn.setEnabled(False)
        for editor in self.coding_editors:
            editor.setReadOnly(True)
            editor.setStyleSheet(editor.styleSheet() + "background-color: #1e2127;")

        self.btn_submit.setText("QUAY LẠI TRANG CHỦ")
        self.btn_submit.setStyleSheet("height: 55px; background-color: #2ecc71; font-size: 16px;")
        self.btn_export.setEnabled(True)
        self.btn_export.setStyleSheet("height: 55px; background-color: #3498db; font-size: 15px; border-radius: 8px;")
        
        correct = sum(1 for i, g in enumerate(self.answer_groups) if g.checkedId() == self.questions[i]['ans'])
        score = (correct / len(self.questions)) * 10 if self.questions else 0
        
        self.generate_document(save_locally=False)
        
        student_answers = [g.checkedId() for g in self.answer_groups]

        res_box = QMessageBox(self)
        res_box.setWindowTitle("Kết quả bài thi")
        status_msg = "Thời gian đã hết!" if auto else "Nộp bài thành công!"
        
        res_box.setText(f"<div style='text-align: center;'><b style='font-size: 18px; color: #e74c3c;'>{status_msg}</b><br>"
                        f"<b style='font-size: 20px; color: #2980b9;'>Học sinh: {self.student_name}</b></div>")
        
        res_box.setInformativeText(f"""
            <div style='text-align: center; font-size: 16px; color: #2f3640;'>
                <p style='font-size: 22px;'>Điểm trắc nghiệm: <span style='color: #e67e22; font-weight: bold;'>{score:.2f}/10</span></p>
                <p style='color: #7f8c8d; font-size: 13px;'>Em có thể xem lại các câu đúng/sai bằng nút bên dưới.</p>
            </div>
        """)
        
        btn_review = res_box.addButton("XEM ĐÁP ÁN", QMessageBox.ActionRole)
        btn_review.setStyleSheet("background-color: #9b59b6; color: white; padding: 10px; font-weight: bold;")
        btn_close = res_box.addButton("ĐÓNG", QMessageBox.AcceptRole)
        btn_close.setStyleSheet("background-color: #3498db; color: white; padding: 10px; font-weight: bold;")
        
        res_box.exec()
        
        if res_box.clickedButton() == btn_review:
            from review import ReviewDialog
            review_window = ReviewDialog(self.questions, student_answers, self)
            review_window.exec()
            

    def generate_document(self, save_locally=False):
        SERVER_IP = "192.168.1.60"
        SHARE_NAME = "KET_QUA_THI"
        NETWORK_PATH = f"//{SERVER_IP}/{SHARE_NAME}"
        file_name = f"{self.quiz_title}_{self.student_name.replace(' ', '_')}.docx"
        
        doc = Document()
        doc.add_heading(f'BÀI KIỂM TRA: {self.quiz_title}', 0)
        doc.add_paragraph(f'Học sinh: {self.student_name}')
        
        doc.add_heading('1. Trắc nghiệm', level=1)
        for i, q in enumerate(self.questions):
            sel = self.answer_groups[i].checkedId()
            ans_text = q['opts'][sel] if sel != -1 else "Bỏ trống"
            result = "ĐÚNG" if sel == q['ans'] else "SAI"
            doc.add_paragraph(f"Câu {i+1}: {q['q']}\n- Trả lời: {ans_text} ({result})\n- Đáp án đúng: {q['opts'][q['ans']]}")

        doc.add_heading('2. Tự luận', level=1)
        for i, editor in enumerate(self.coding_editors):
            doc.add_heading(f'Bài tập {i+1}', level=2)
            doc.add_paragraph(editor.toPlainText())
        
        if save_locally:
            try:
                doc.save(file_name)
                self.show_custom_info("Thành công", f"Đã lưu bản sao tại: {file_name}")
            except Exception as e: QMessageBox.critical(self, "Lỗi", str(e))
        else:
            try:
                if os.path.exists(NETWORK_PATH):
                    doc.save(os.path.join(NETWORK_PATH, file_name))
                else: doc.save(file_name)
            except: pass

    def show_custom_info(self, title, text):
        m = QMessageBox(self)
        m.setWindowTitle(title)
        m.setText(f"<span style='color: #2f3640;'>{text}</span>")
        m.exec()
        
    def closeEvent(self, event):
        """Xử lý sự kiện khi người dùng nhấn nút X hoặc Alt+F4"""
        
        # Nếu đã nộp bài rồi thì cho phép thoát luôn không cần hỏi
        if self.is_submitted:
            event.accept()
            return

        # Nếu chưa nộp bài, hiển thị bảng cảnh báo
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Cảnh báo thoát ứng dụng")
        msg_box.setIcon(QMessageBox.Warning)
        msg_box.setText("<b style='font-size: 15px; color: #c0392b;'>EM CÓ CHẮC CHẮN MUỐN THOÁT KHÔNG?</b>")
        msg_box.setInformativeText("Nếu thoát bây giờ, toàn bộ bài làm của em sẽ bị mất và không được ghi nhận.")
        
        btn_exit = msg_box.addButton("Thoát (Hủy bài)", QMessageBox.DestructiveRole)
        btn_stay = msg_box.addButton("Quay lại làm tiếp", QMessageBox.AcceptRole)
        
        btn_exit.setStyleSheet("background-color: #e74c3c; color: white; font-weight: bold; min-width: 100px;")
        btn_stay.setStyleSheet("background-color: #27ae60; color: white; font-weight: bold; min-width: 100px;")
        
        msg_box.exec()
        
        if msg_box.clickedButton() == btn_stay:
            event.ignore()
        else:
            event.accept()

if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setWindowIcon(QIcon(resource_path("logo_cnet.ico")))
    window = QuizApp()
    window.show()
    sys.exit(app.exec())