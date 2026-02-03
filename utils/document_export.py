import os
from docx import Document
from config import SERVER_IP, SHARE_NAME, NETWORK_PATH


class DocumentExporter:
    
    @staticmethod
    def generate_document(quiz_title, student_name, questions, student_answers, coding_tasks, coding_contents):
        doc = Document()

        doc.add_heading(f'BÀI KIỂM TRA: {quiz_title}', 0)
        doc.add_paragraph(f'Học sinh: {student_name}')
        
        doc.add_heading('1. Trắc nghiệm', level=1)
        for i, q in enumerate(questions):
            sel = student_answers[i]
            ans_text = q['opts'][sel] if sel != -1 else "Bỏ trống"
            result = "ĐÚNG" if sel == q['ans'] else "SAI"
            doc.add_paragraph(
                f"Câu {i+1}: {q['q']}\n"
                f"- Trả lời: {ans_text} ({result})\n"
                f"- Đáp án đúng: {q['opts'][q['ans']]}"
            )
        
        # Coding section
        doc.add_heading('2. Tự luận', level=1)
        for i, content in enumerate(coding_contents):
            doc.add_heading(f'Bài tập {i+1}', level=2)
            doc.add_paragraph(content)
        
        return doc
    
    @staticmethod
    def save_document_locally(doc, quiz_title, student_name):
        #Save document locally (Teacher's computer)
        file_name = f"{quiz_title}_{student_name.replace(' ', '_')}.docx"
        try:
            doc.save(file_name)
            return True, file_name
        except Exception as e:
            return False, str(e)
    
    @staticmethod
    def save_document_to_network(doc, quiz_title, student_name):
        #Save document to network share (Student's computer)
        file_name = f"{quiz_title}_{student_name.replace(' ', '_')}.docx"
        try:
            if os.path.exists(NETWORK_PATH):
                doc.save(os.path.join(NETWORK_PATH, file_name))
                return True, file_name
            else:
                return DocumentExporter.save_document_locally(doc, quiz_title, student_name)
        except Exception:
            return DocumentExporter.save_document_locally(doc, quiz_title, student_name)
